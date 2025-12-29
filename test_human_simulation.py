import sys
import os
import time
from logic.role_workflows import CDMWorkflows, CRAWorkflows, QAWorkflows, OncologyWorkflows
from logic.agent_logic import generate_protocol_draft

# Mock File Object for Uploads
class MockFile:
    def __init__(self, name):
        self.name = name

# Helper
def auto_open(fpath):
    # Check for CI/Headless Mode (Set Env Var CI=true to skip popups)
    if os.environ.get("CI") == "true":
        print(f"🔒 CI MODE: Skipping open for {os.path.basename(fpath)}")
        return

    if os.name == 'nt' and os.path.exists(fpath):
        print(f"📂 POP-UP: Opening {os.path.basename(fpath)}...")
        os.startfile(os.path.abspath(fpath))

print("🤖 STARTING FULL 'HUMAN' SIMULATION...")
print("=" * 60)

# 1. MEDICAL WRITER: Protocol Draft
print("\n[1] Role: Medical Writer")
print("👉 Action: Generate 'Novartis-Style' Protocol Draft (DOCX)...")
try:
    draft = generate_protocol_draft("Phase 3", "RCT", "Study of Drug X in Asthma", "Focus on safety")
    fname = "Protocol_Draft_SIM.docx"
    
    # Verify Word Logic
    from docx import Document
    doc = Document()
    doc.add_heading('Clinical Trial Protocol', 0)
    for line in draft.split('\n'):
        if line.strip():
            doc.add_paragraph(line)
    doc.save(fname)
    
    if "Novartis" not in draft and "[SPONSOR]" in draft:
        print("✅ SUCCESS: Protocol Generated & De-Branded (DOCX Saved).")
        auto_open(fname)
    else:
        print("⚠️ CHECK: Branding may persist.")
except Exception as e:
    print(f"❌ FAIL: {e}")

# 2. QA AUDITOR: eTMF Filing
print("\n[2] Role: Clinical QA / Auditor")
print("👉 Action: File a Lab Report...")
try:
    mock_file = MockFile("Site_101_Lab_Results.pdf")
    res = QAWorkflows.audit_and_file_document(mock_file)
    if res["Status"] == "Filed":
        print(f"✅ SUCCESS: {res['Message']}")
        print(f"   (Mapped to Zone: {res['Zone']})")
    else:
        print(f"❌ FAIL: {res['Message']}")
except Exception as e:
    print(f"❌ FAIL: {e}")

# 3. ONCOLOGY: RECIST Calc
print("\n[3] Role: Medical Monitor (Oncology)")
print("👉 Action: Calculate Tumor Response (Base=100, Nadir=60, Curr=80)...")
try:
    # 100 -> 60 (Nadir) -> 80 (+20mm, +33%). Should be PD (>20% from Nadir).
    recist = OncologyWorkflows.calculate_recist(100, 60, 80)
    print(f"   Result: {recist['Response']}")
    if "PD" in recist["Response"]:
        print("✅ SUCCESS: Correctly identified Progressive Disease (PD).")
    else:
        print(f"❌ FAIL: Expected PD, got {recist['Response']}")
except Exception as e:
    print(f"❌ FAIL: {e}")

# 4. CRA: Trip Report
print("\n[4] Role: CRA / Monitor")
try:
    pdf_path = CRAWorkflows.generate_trip_report("Site visit ok. Drug logs compliant.")
    if "Error" not in pdf_path:
        print("✅ SUCCESS: Trip Report PDF Generated.")
        auto_open(pdf_path)
except Exception as e:
    print(f"❌ FAIL: {e}")

print("=" * 60)
print("🏁 SIMULATION COMPLETE.")
