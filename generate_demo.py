import os
import time
from docx import Document
from logic.agent_logic import generate_protocol_draft

# Complex Scenario to test "Intelligence"
PHASE = "Phase 3"
DESIGN = "Randomized, Double-Blind, Placebo-Controlled, Parallel-Group"
TITLE = "A Phase 3 Study to Evaluate Efficacy and Safety of Neuro-X in Early Alzheimer's Disease"
CONTEXT = """
Target: Early AD patients (MCI). 
Drug mechanism: Amyloid-beta clearance. 
Primary Endpoint: Change in CDR-SB at Week 78.
Secondary: ADAS-Cog13, ADCS-ADL.
Safety: ARIA-E monitoring via MRI.
"""

print(f"🚀 GENERATING PROTOCOL ZERO: {TITLE}...")
print("(This uses the new Mega-Prompt for 80-90% completion...)")

try:
    draft = generate_protocol_draft(PHASE, DESIGN, TITLE, CONTEXT)
    
    fname = f"DEMO_Protocol_Zero_{int(time.time())}.docx"
    
    # Word Conversion Logic
    doc = Document()
    doc.add_heading('Clinical Trial Protocol (Zero Draft)', 0)
    
    for line in draft.split('\n'):
        if line.strip():
            if line.strip().startswith('#'):
                 clean_head = line.strip().replace('#', '').strip()
                 # Simple heuristic for levels
                 level = 1 if line.count('#') == 1 else 2
                 doc.add_heading(clean_head, level=1) # Keep it simple for demo
            else:
                doc.add_paragraph(line)
                
    doc.save(fname)
    print(f"✅ GENERATED: {os.path.abspath(fname)}")
    
    # Open for User
    os.startfile(os.path.abspath(fname))
    
except Exception as e:
    print(f"❌ ERROR: {e}")
